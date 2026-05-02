"""
Word 文档导出模块

将转录文本整理为格式化的 .docx 文件。

用法:
    python -m src.exporter
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config import load_config


def clean_text(raw: str) -> str:
    """去除时间戳，合并为连续文本"""
    lines = []
    for line in raw.strip().split("\n"):
        cleaned = re.sub(r"\[[\d:]+\s*→\s*[\d:]+\]\s*", "", line).strip()
        if cleaned:
            lines.append(cleaned)
    return "".join(lines)


def text_to_paragraphs(text: str, sentences_per_para: int = 4) -> list[str]:
    """按句号拆分为段落"""
    sentences = re.split(r"(?<=[。！？])", text)
    sentences = [s.strip() for s in sentences if s.strip()]

    paragraphs = []
    buf = []
    for sent in sentences:
        buf.append(sent)
        if len(buf) >= sentences_per_para:
            paragraphs.append("".join(buf))
            buf = []
    if buf:
        paragraphs.append("".join(buf))
    return paragraphs


def export_one(txt_path: Path, word_dir: Path) -> Path | None:
    """将单个 txt 转为 Word"""
    raw = txt_path.read_text(encoding="utf-8")
    if not raw.strip():
        print(f"  [跳过] {txt_path.name}（空文件）")
        return None

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(12)

    # 解析标题：如果文件名格式是 "博主名_时间戳"，分别提取
    stem = txt_path.stem
    if "_" in stem:
        # 尝试拆分为 博主名 和 时间戳
        parts = stem.split("_", 1)
        creator_name = parts[0]
        timestamp_part = parts[1] if len(parts) > 1 else ""
        title = f"{creator_name}"
        subtitle = timestamp_part.replace("_", " ")
    else:
        title = stem.replace("_", " ")
        creator_name = ""
        subtitle = ""

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = f"转录日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}  |  工具：Douyin-to-Text"
    if subtitle:
        meta_text = f"录制时间：{subtitle}  |  {meta_text}"
    run = meta.add_run(meta_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("")

    clean = clean_text(raw)
    paragraphs = text_to_paragraphs(clean)

    for para_text in paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(para_text)
        r.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Pt(24)

    word_path = word_dir / f"{txt_path.stem}.docx"
    doc.save(str(word_path))
    return word_path


def run():
    """导出所有未处理的转录文本"""
    config = load_config()
    transcript_dir = Path(config["paths"]["transcript_dir"])
    word_dir = Path(config["paths"]["word_dir"])

    txt_files = sorted(transcript_dir.glob("*.txt"))
    if not txt_files:
        print(f"[提示] {transcript_dir} 下没有待导出文本")
        return

    exported = 0
    for tf in txt_files:
        word_path = word_dir / f"{tf.stem}.docx"
        if word_path.exists():
            print(f"[跳过] {tf.stem}（Word 已存在）")
            continue

        print(f"[导出] {tf.name}")
        result = export_one(tf, word_dir)
        if result:
            print(f"  → {result}")
            exported += 1

    print(f"\n[完成] 导出 {exported} 个 Word 到 {word_dir}")


if __name__ == "__main__":
    run()
