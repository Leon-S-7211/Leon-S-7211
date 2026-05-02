"""
DeepSeek 大模型分析模块

根据博主的 category 配置（finance / career）自动选择对应的提示词：
- finance → prompt_finance.txt（金融分析）
- career  → prompt_career.txt（求职分析）

分析结果保存到 "文献" 文件夹中。

用法:
    python -m src.analyzer                 # 分析所有未处理的 Word
    python -m src.analyzer --file xxx.docx # 分析指定文件

配置:
    config.json 中每个博主需要有 "category" 字段:
    { "name": "机构一手调研", "category": "finance", ... }
    { "name": "赫子", "category": "career", ... }
"""

import sys
from pathlib import Path
from datetime import datetime

from .config import load_config

# 提示词文件路径
PROMPT_DIR = Path(__file__).parent.parent
PROMPT_FILES = {
    "finance": PROMPT_DIR / "prompt_finance.txt",
    "career": PROMPT_DIR / "prompt_career.txt",
}

# 类型的中文标签（用于文件名和报告标题）
CATEGORY_LABELS = {
    "finance": "金融分析",
    "career": "求职分析",
}


def build_creator_category_map(config: dict) -> dict[str, str]:
    """
    从 config.json 构建 博主名 → category 的映射表。

    同时读取 creators 和 live_creators 两个列表。
    """
    mapping = {}
    for creator in config.get("creators", []):
        name = creator.get("name", "")
        cat = creator.get("category", "")
        if name and cat:
            mapping[name] = cat

    for creator in config.get("live_creators", []):
        name = creator.get("name", "")
        cat = creator.get("category", "")
        if name and cat:
            mapping[name] = cat

    return mapping


def match_category(filename: str, category_map: dict) -> str | None:
    """
    从 Word 文件名中匹配博主名，找到对应的 category。

    文件名格式: "博主名_2026-04-27 10-47-31.docx"
    所以取下划线前面的部分跟 category_map 的 key 匹配。
    """
    stem = Path(filename).stem

    # 直接匹配文件名开头
    for creator_name, category in category_map.items():
        if stem.startswith(creator_name):
            return category

    # 如果文件名里包含博主名（兼容各种命名格式）
    for creator_name, category in category_map.items():
        if creator_name in stem:
            return category

    return None


def load_prompt(category: str) -> str:
    """加载指定类型的提示词文件"""
    prompt_file = PROMPT_FILES.get(category)

    if not prompt_file:
        print(f"[错误] 未知的 category: {category}")
        print(f"       支持的值: {', '.join(PROMPT_FILES.keys())}")
        return ""

    if not prompt_file.exists():
        print(f"[错误] 找不到提示词文件: {prompt_file}")
        print(f"       请创建该文件并写入 {CATEGORY_LABELS.get(category, category)} 的提示词。")
        return ""

    text = prompt_file.read_text(encoding="utf-8").strip()
    if not text:
        print(f"[错误] {prompt_file.name} 是空的。")
        return ""

    return text


def extract_text_from_docx(docx_path: Path) -> str:
    """从 Word 文档中提取纯文本"""
    from docx import Document
    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def call_deepseek(config: dict, system_prompt: str, content: str) -> str:
    """调用 DeepSeek API"""
    try:
        from openai import OpenAI
    except ImportError:
        print("[错误] 请安装 openai 库: pip install openai")
        sys.exit(1)

    ds_config = config.get("deepseek", {})
    api_key = ds_config.get("api_key", "")
    model = ds_config.get("model", "deepseek-chat")
    base_url = ds_config.get("base_url", "https://api.deepseek.com")

    if not api_key:
        print("[错误] 请在 config.json 中配置 deepseek.api_key")
        sys.exit(1)

    client = OpenAI(api_key=api_key, base_url=base_url)

    print("  [API] 正在调用 DeepSeek...")

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": content},
            ],
            max_tokens=4096,
            temperature=0.3,
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"  [错误] DeepSeek API 调用失败: {e}")
        return ""


def save_analysis(analysis: str, source_name: str, category: str, output_dir: Path) -> Path:
    """将分析结果保存为 Word 文档"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    label = CATEGORY_LABELS.get(category, "分析")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(12)

    heading = doc.add_heading(f"{label}报告：{source_name}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"生成日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"类型：{label}  |  模型：DeepSeek"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("")

    for para in analysis.split("\n"):
        para = para.strip()
        if not para:
            doc.add_paragraph("")
            continue

        if para.startswith("### "):
            doc.add_heading(para[4:], level=4)
        elif para.startswith("## "):
            doc.add_heading(para[3:], level=3)
        elif para.startswith("# "):
            doc.add_heading(para[2:], level=2)
        elif para.startswith("【") and para.endswith("】"):
            doc.add_heading(para, level=2)
        else:
            p = doc.add_paragraph()
            r = p.add_run(para)
            r.font.size = Pt(12)
            p.paragraph_format.line_spacing = 1.5

    output_name = f"{source_name}_{label}.docx"
    output_path = output_dir / output_name
    doc.save(str(output_path))
    return output_path


def run(target_file: str = None):
    """分析所有未处理的转录 Word"""
    config = load_config()
    word_dir = Path(config["paths"]["word_dir"])

    analysis_dir = Path(config["paths"].get(
        "analysis_dir",
        str(Path(config["paths"]["word_dir"]).parent / "文献")
    ))
    analysis_dir.mkdir(parents=True, exist_ok=True)

    # 构建博主→category 映射
    category_map = build_creator_category_map(config)
    if category_map:
        print("[配置] 博主类型映射:")
        for name, cat in category_map.items():
            print(f"  • {name} → {CATEGORY_LABELS.get(cat, cat)}")
        print()
    else:
        print("[警告] config.json 中没有博主配置了 category 字段。")
        print("       请给每个博主添加 \"category\": \"finance\" 或 \"career\"")
        return

    # 确定要分析的文件
    if target_file:
        docx_files = [Path(target_file)]
    else:
        docx_files = sorted(word_dir.glob("*.docx"))

    if not docx_files:
        print(f"[提示] {word_dir} 下没有 Word 文件")
        return

    analyzed = 0
    skipped_no_category = []

    for docx_path in docx_files:
        source_name = docx_path.stem

        # 匹配 category
        category = match_category(docx_path.name, category_map)
        if not category:
            skipped_no_category.append(docx_path.name)
            continue

        label = CATEGORY_LABELS.get(category, "分析")

        # 跳过已有分析报告的
        report_path = analysis_dir / f"{source_name}_{label}.docx"
        if report_path.exists():
            print(f"[跳过] {source_name}（{label}报告已存在）")
            continue

        print(f"[分析] {docx_path.name}  →  {label}")

        # 加载提示词
        prompt = load_prompt(category)
        if not prompt:
            continue

        # 提取文本
        text = extract_text_from_docx(docx_path)
        if not text.strip():
            print(f"  [跳过] 内容为空")
            continue

        print(f"  [内容] {len(text)} 字")

        # 调 DeepSeek
        analysis = call_deepseek(config, prompt, text)
        if not analysis:
            print(f"  [跳过] 分析结果为空")
            continue

        # 保存
        result_path = save_analysis(analysis, source_name, category, analysis_dir)
        print(f"  [保存] {result_path}")
        analyzed += 1

    # 提示未匹配的文件
    if skipped_no_category:
        print(f"\n[提示] 以下文件无法匹配博主 category，跳过分析:")
        for f in skipped_no_category:
            print(f"  • {f}")
        print("       请检查文件名是否包含博主名，或在 config.json 中添加 category。")

    print(f"\n[完成] 分析 {analyzed} 个文档，结果在: {analysis_dir}")


if __name__ == "__main__":
    target = None
    if len(sys.argv) > 2 and sys.argv[1] == "--file":
        target = sys.argv[2]
    run(target_file=target)
